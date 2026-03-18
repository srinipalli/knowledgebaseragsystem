from __future__ import annotations

from pathlib import Path

from docx import Document

from app.parsers.base import BaseParser, TextUnit


class DOCXParser(BaseParser):
    name = "python-docx"

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def estimate_total_units(self, path: Path) -> int:
        document = Document(path)
        return len(document.paragraphs)

    def iter_units(self, path: Path):
        document = Document(path)
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            yield TextUnit(
                unit_index=index,
                text=text,
                metadata={"paragraph_number": index, "source_type": "docx"},
            )
